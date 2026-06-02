from __future__ import annotations

import ast
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]


APPENDIX_TITLE = "Appendix A. Variable Range Interpretation Guide"

SLIDE_MAP = [
    ("Future Well-Log Engine boundary banner", "Slide 1 and Slide 12", "Public-source scaffold boundary and authorized-runtime transfer rule"),
    ("Scientific rules kept visible", "Slide 2", "Decision stack: admissibility, reservoir, phase, saturation, and producibility remain separate"),
    ("Synthetic runtime schema", "Slide 3", "Curve-to-purpose map and runtime configuration layer"),
    ("Hydrate Interpretation Range Guide", "Slide 4 and Slide 11", "Physics-backed feature tendencies, competing explanations, and safeguards"),
    ("Variable Range Explorer", "Slide 6 and Slide 9", "Per-curve QC statistics, depth panel, and export-ready comparison table"),
    ("Interval Screening Scaffold", "Slide 7 and Slide 9", "Separate staged outputs including good-sand/no-hydrate and expert-review outcomes"),
    ("Core Calibration Scaffold", "Slide 9", "Depth-match window, confidence weight, disturbance flag, and nearby-log linkage"),
    ("Presentation Outputs", "Slide 10", "Synthetic confusion matrix, calibration panel, uncertainty summary, and downloadable placeholders"),
]


def load_range_guide() -> list[dict[str, str]]:
    source = (REPO_ROOT / "dashboard" / "well_log_engine.py").read_text(encoding="utf-8")
    module = ast.parse(source)
    for node in module.body:
        if isinstance(node, ast.Assign) and any(
            isinstance(target, ast.Name) and target.id == "RANGE_GUIDE"
            for target in node.targets
        ):
            return ast.literal_eval(node.value)
    raise RuntimeError("RANGE_GUIDE was not found in dashboard/well_log_engine.py")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.5, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C9D3")
        borders.append(element)
    tbl_pr.append(borders)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_reference_table(doc: Document, range_guide: list[dict[str, str]]) -> None:
    headers = ["Variable", "Unit", "Working tendency", "Why it can help", "Competing explanations / required evidence", "Uncertainty warning", "Manuscript source"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_widths(table, [0.70, 0.62, 1.45, 1.50, 2.30, 1.65, 1.55])
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "123447")
        set_cell_text(cell, header, bold=True, size=7.3, color="FFFFFF")
    for item in range_guide:
        row = table.add_row().cells
        values = [
            item["Variable"],
            item["Unit"],
            item["Working tendency or screening range"],
            item["Why it may support interpretation"],
            f'{item["Competing explanations"]} Required evidence: {item["Required supporting evidence"]}',
            item["Uncertainty warning"],
            item["Manuscript source"],
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value)


def add_slide_map(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_widths(table, [2.75, 1.35, 5.45])
    for cell, header in zip(table.rows[0].cells, ["Dashboard output", "Planned slide(s)", "Presentation use"]):
        shade_cell(cell, "167D8D")
        set_cell_text(cell, header, bold=True, size=8, color="FFFFFF")
    for output, slides, use in SLIDE_MAP:
        row = table.add_row().cells
        for cell, value in zip(row, [output, slides, use]):
            set_cell_text(cell, value, size=8)


def add_boundary_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F4EFE6")
    set_cell_text(
        cell,
        "Boundary rule: This appendix and the hosted dashboard are PUBLIC-SOURCE PLANNING SCAFFOLDS. "
        "They use synthetic examples only. Future approved well logs, core data, identifiers, populated "
        "sensitive outputs, derived sensitive results, and credentials must remain inside the authorized "
        "DOE environment and must not be committed to GitHub, uploaded to Streamlit Community Cloud, or placed in chat.",
        bold=True,
        size=9,
    )


def append_blueprint(path: Path) -> None:
    range_guide = load_range_guide()
    doc = Document(path)
    if any(paragraph.text.strip() == APPENDIX_TITLE for paragraph in doc.paragraphs):
        raise SystemExit(f"{APPENDIX_TITLE!r} is already present; refusing to append a duplicate.")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title = doc.add_heading(APPENDIX_TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(
        "This appendix translates the Alaska North Slope manuscript into a dashboard-ready reference. "
        "The listed values are overlapping working anchors and response tendencies, not universal decision thresholds."
    )
    add_boundary_callout(doc)
    doc.add_paragraph()

    doc.add_heading("A.1 Interpretation Rules", level=2)
    for rule in [
        "GHSZ is necessary but not sufficient.",
        "High Rt is evidence, not a hydrate label.",
        "Good reservoir sand can contain no hydrate.",
        "Geology and seismic context constrain confidence; they do not replace direct log evidence.",
        "NMR-density saturation is preferred where NMR exists; Archie is a supplementary cross-check with uncertainty flags.",
        "Hydrate occurrence, saturation, and producibility remain separate outputs.",
        "Validation must split by well, not randomly by depth sample.",
        "Maximum hydrate saturation is not automatically the best production target.",
    ]:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("A.2 Variable Range Interpretation Guide", level=2)
    doc.add_paragraph(
        "Use each tendency in a multi-log interpretation chain. Competing explanations and required supporting evidence "
        "remain visible because the manuscript explicitly preserves false positives and ambiguous intervals."
    )
    add_reference_table(doc, range_guide)

    doc.add_heading("A.3 Multi-Log Confirmation Rules", level=2)
    for rule in [
        "Apply pressure-temperature context as an admissibility screen. Do not use stability as a positive label.",
        "Screen reservoir quality before phase. Preserve reservoir-grade water-bearing sand as a valid outcome.",
        "Treat high Rt as supporting evidence only when lithology, density/porosity, sonic behavior, and borehole QC agree.",
        "Prefer NMR-density saturation where NMR is available. Retain Archie as a supplementary estimate with salinity, lithology, and low-porosity uncertainty flags.",
        "Use Vp, Vs, Vp/Vs, lambda-rho, and mu-rho together. Review stress, ice-bearing sediment, competent lithology, and depth context before assigning a hydrate-supportive interpretation.",
        "Keep core sample depth-match uncertainty, disturbance flags, pressure-core quality, and confidence weights attached to any future calibration label.",
        "Rank producibility separately from occurrence and saturation because retained permeability and pressure communication can decline as saturation rises.",
    ]:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("A.4 Dashboard Outputs to Planned Slides", level=2)
    doc.add_paragraph(
        "The hosted scaffold exposes synthetic examples for presentation planning. Approved-data replacements must be generated locally inside the authorized environment."
    )
    add_slide_map(doc)
    doc.save(path)


if __name__ == "__main__":
    append_blueprint(Path(sys.argv[1]))
