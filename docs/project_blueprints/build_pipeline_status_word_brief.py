from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "PIPELINE_STATUS_AND_ML_WORKFLOW_BRIEF.md"
OUT_DOCX = (
    ROOT
    / "docs"
    / "project_blueprints"
    / "North_Slope_Gas_Hydrate_ML_Pipeline_Status_And_Forward_Workflow_2026-06-15.docx"
)


def clean_cell(value: str) -> str:
    value = value.strip()
    if value.startswith("`") and value.endswith("`") and len(value) > 1:
        value = value[1:-1]
    return value.replace("<br>", "\n")


def is_table_separator(line: str) -> bool:
    text = line.strip()
    if not (text.startswith("|") and text.endswith("|")):
        return False
    parts = [part.strip() for part in text.strip("|").split("|")]
    return bool(parts) and all(part and set(part) <= {"-", ":"} for part in parts)


def table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i]
        if not line.strip().startswith("|"):
            break
        if not is_table_separator(line):
            rows.append([clean_cell(part) for part in line.strip().strip("|").split("|")])
        i += 1
    return rows, i


def apply_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [
        ("Title", 22),
        ("Heading 1", 17),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Title" or True

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_metadata(document: Document) -> None:
    props = document.core_properties
    props.title = "North Slope Gas Hydrate ML Pipeline Status And Forward Workflow"
    props.subject = "Current project status and source-backed ML pipeline plan"
    props.author = "North Slope Gas Hydrates project"
    props.keywords = "gas hydrate, North Slope, stability screen, machine learning, well logs"


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    document.add_paragraph()


def add_code_block(document: Document, block: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run("\n".join(block))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    in_code = False
    code_block: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_block)
                code_block = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_block.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = table_rows(lines, i)
            add_table(document, rows)
            i = next_i
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:4]:
            document.add_paragraph(stripped.split(". ", 1)[1].strip(), style="List Number")
        else:
            document.add_paragraph(stripped)

        i += 1

    if code_block:
        add_code_block(document, code_block)


def build_docx() -> Path:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)

    document = Document()
    apply_document_style(document)
    add_metadata(document)
    add_markdown(document, SOURCE_MD.read_text(encoding="utf-8"))

    # Keep a short appendix marker if reviewers add comments in Word later.
    document.add_section(WD_SECTION.CONTINUOUS)
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Review Note", level=1)
    document.add_paragraph(
        "This document is a public-safe planning brief. It does not contain approved "
        "well-log rows, core rows, restricted identifiers, trained model results, "
        "or final hydrate occurrence/saturation claims."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)
    return OUT_DOCX


def main() -> None:
    path = build_docx()
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
