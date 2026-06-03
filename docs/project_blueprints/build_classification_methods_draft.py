from __future__ import annotations

import ast
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
OUT = REPO_ROOT / "docs" / "project_blueprints" / "Alaska_North_Slope_Wireline_ML_Classification_Methods_Draft.docx"

NAVY = "123447"
TEAL = "167D8D"
ICE = "EAF5F6"
SAND = "F4EFE6"
GRAY = "F2F4F7"
INK = "1F2933"


def load_engine_constant(name: str):
    source = (REPO_ROOT / "dashboard" / "well_log_engine.py").read_text(encoding="utf-8")
    module = ast.parse(source)
    for node in module.body:
        if isinstance(node, ast.Assign) and any(
            isinstance(target, ast.Name) and target.id == name for target in node.targets
        ):
            return ast.literal_eval(node.value)
    raise RuntimeError(f"{name} was not found in dashboard/well_log_engine.py")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_borders(table, color="B7C9D3") -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tbl_pr.append(borders)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row.cells[index])


def write_cell(cell, text: str, *, bold=False, size=8.4, color=INK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=TEAL) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_borders(table)
    set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, header_fill)
        write_cell(cell, header, bold=True, color="FFFFFF", size=8)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            write_cell(cell, value)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str, fill=SAND) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_borders(table, color="D2C8B8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else TEAL)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(10)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def main() -> None:
    equation_library = load_engine_constant("EQUATION_LIBRARY")
    range_guide = load_engine_constant("RANGE_GUIDE")
    rocktype_context_guide = load_engine_constant("ROCKTYPE_CONTEXT_GUIDE")
    classification_workflow = load_engine_constant("CLASSIFICATION_WORKFLOW")
    sweet_spot_guide = load_engine_constant("SWEET_SPOT_GUIDE")

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Alaska North Slope Wireline ML Classification Methods Draft")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph(
        "Unclassified scaffold for future approved well-log and core data: variables to equations to staged classification to ML outputs"
    )
    subtitle.paragraph_format.space_after = Pt(10)
    for run in subtitle.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string("556B73")

    add_callout(
        doc,
        "Boundary",
        "This document is a public-source planning and methods scaffold. It contains no classified or restricted well logs, core measurements, identifiers, populated results, or derived sensitive outputs. Future approved data must remain in the authorized runtime environment.",
    )

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "This draft reframes the Alaska North Slope gas hydrate project as a physics-constrained wireline machine-learning workflow. The purpose is not to restate general gas hydrate geology, but to define how future approved well logs and core observations will be converted into interpretable features, staged hydrate evidence, uncertainty labels, producibility screens, and presentation-ready results. The working hypothesis is that hydrate classification requires six separable decisions: data quality, pressure-temperature admissibility, reservoir quality, multi-log phase evidence, competing rock-type explanations, and producibility risk. The scaffold therefore treats every sweet spot as conditional: a variable range is useful only when the supporting curves, rock type, overburden/stress context, and calibration evidence agree.",
    )

    add_heading(doc, "1. Project Direction", 1)
    add_body(
        doc,
        "The earlier manuscript is best treated as a source accumulation and scientific synthesis. This new direction is a methods paper and dashboard specification for a future real-data project. Its central chain is: measured curves -> derived equations -> physical variables -> staged interpretation gates -> ML features -> interval classifications -> results and discussion.",
    )
    add_bullets(
        doc,
        [
            "Primary task: classify hydrate-supportive, gas-supportive, water-bearing, non-reservoir, and expert-review intervals from approved wireline and core data.",
            "Secondary task: rank candidate sweet spots without confusing hydrate detectability, hydrate saturation, and producibility.",
            "Scientific control: every range must carry its physical reason, competing explanations, and required supporting evidence.",
            "ML control: features should be grouped by well and geology so validation does not leak depth-adjacent samples between training and testing.",
        ],
    )

    add_heading(doc, "2. Measurement Families and Runtime Inputs", 1)
    measurement_rows = [
        ["Lithology", "GR, optional mineralogy/core", "Vsh, clean-sand flag, rock-type screen", "Defines whether a positive response is in plausible reservoir rock."],
        ["Electrical", "Rt, Rw/calibration if available", "Archie hydrate proxy", "Hydrate can raise resistivity by reducing connected brine pathways, but gas, ice, coal, and low porosity can also raise Rt."],
        ["Density/porosity", "RHOB, NPHI, matrix/fluid assumptions", "Density porosity, storage capacity, density QC", "Separates pore-volume capacity from hydrate occupancy."],
        ["Sonic/elastic", "DT, DTS, RHOB", "Vp, Vs, Vp/Vs, G, K, lambda-rho, mu-rho, Poisson ratio, Young's modulus", "Tests whether high Rt is accompanied by hydrate/ice/competent-rock stiffness or gas-like softening."],
        ["NMR/core", "NMR porosity, core porosity/saturation, pressure-core quality", "NMR-density hydrate proxy, confidence weights", "Preferred calibration path where available; preserves depth-match uncertainty."],
        ["Thermal/stress", "Depth, temperature gradient, pressure, overburden gradient", "GHSZ admissibility, vertical stress, effective stress", "Hydrate must be stable, and stress changes velocity/porosity interpretation."],
    ]
    add_table(doc, ["Family", "Measured inputs", "Derived features", "Decision role"], measurement_rows, [1.05, 1.55, 2.35, 2.35])

    add_heading(doc, "3. Equation Library", 1)
    add_body(
        doc,
        "The equation library below is the minimum mathematical scaffold to represent the future wellbore workflow. The equations are not final field calibrations; they are traceable transforms that define what the dashboard and model should calculate once approved data are available.",
    )
    equation_rows = [
        [row["Equation group"], row["Equation"], row["Feature produced"], row["Classification use"]]
        for row in equation_library
    ]
    add_table(doc, ["Equation group", "Equation", "Feature", "How it enters classification"], equation_rows, [1.35, 2.35, 1.55, 2.45], header_fill=NAVY)

    add_heading(doc, "4. Why Sweet Spots Are Conditional", 1)
    add_body(
        doc,
        "A sweet spot is not a single threshold. For this project, a sweet spot is an interval that passes a staged set of physical screens: admissible pressure-temperature conditions, reservoir-quality rock, multi-log hydrate evidence, limited competing explanations, and enough retained flow capacity to justify production-oriented review. The same resistivity or velocity value can mean different things in clean sand, shaly sand, permafrost/ice-bearing sediment, coal, carbonate, or a washed-out borehole.",
    )
    range_rows = [
        [
            item["Variable"],
            item["Working tendency or screening range"],
            item["Why it may support interpretation"],
            item["Competing explanations"],
            item["Required supporting evidence"],
        ]
        for item in range_guide
        if item["Variable"] in {"GR", "Rt", "Porosity", "Vp", "Vs", "lambda-rho", "mu-rho", "NMR-density separation", "Temperature", "Pressure"}
    ]
    add_table(doc, ["Variable", "Working tendency", "Why useful", "Competing explanations", "Required support"], range_rows, [0.75, 1.75, 1.75, 1.85, 1.85], header_fill=TEAL)

    add_heading(doc, "5. Overburden, Pressure, and Rock-Type Effects", 1)
    add_body(
        doc,
        "Overburden and pore pressure affect the interpretation because increasing effective stress can compact sediments, reduce porosity, and increase velocity even without hydrate. Conversely, high pore pressure can reduce effective stress and alter mechanical response. The dashboard therefore calculates stress context as a review variable, not as a hydrate detector. Rock type is treated the same way: it changes the meaning of each curve and must be attached to every classification lane.",
    )
    rock_rows = [
        [row["Rock / interval type"], row["Why it helps"], row["False-positive risk"], row["Dashboard response"]]
        for row in rocktype_context_guide
    ]
    add_table(doc, ["Rock/context", "Why it matters", "False-positive risk", "Model/dashboard response"], rock_rows, [1.35, 2.0, 2.05, 2.05], header_fill=NAVY)

    add_heading(doc, "6. Staged Classification Workflow", 1)
    add_body(
        doc,
        "The classifier should act like a disciplined interpreter. It should not jump from a curve value to a label. It should carry an interval through gates, keep uncertainty visible, and preserve good-sand/no-hydrate and expert-review outcomes as valid scientific results.",
    )
    workflow_rows = [
        [row["Stage"], row["Question"], row["Primary variables"], row["Pass logic"]]
        for row in classification_workflow
    ]
    add_table(doc, ["Stage", "Question", "Primary variables", "Pass logic"], workflow_rows, [1.2, 2.15, 2.05, 2.3], header_fill=TEAL)

    add_heading(doc, "7. Machine-Learning Design", 1)
    add_body(
        doc,
        "The ML system should begin with interpretable baselines and only move toward more flexible models after grouped validation is stable. The first defensible model ladder is: rule-based baseline, regularized logistic regression or calibrated random forest, gradient-boosted trees, and then sequence-aware or interval-aware models if the data volume and labels justify them. The model should output probabilities and reasons, not just class labels.",
    )
    add_bullets(
        doc,
        [
            "Features: raw curves, derived equations, QC flags, rock-type screen, pressure-temperature context, overburden/effective stress, and interval statistics.",
            "Labels: hydrate-supportive, gas-supportive, water sand, good sand/no hydrate, non-reservoir, and expert review.",
            "Validation: split by well or fault/reservoir compartment, not by random depth sample.",
            "Calibration: compare probability calibration, confusion matrix, feature importance, and failure cases against core-calibrated intervals.",
            "Outputs: interval table, depth-panel callouts, cross-well variable ranges, sweet-spot lanes, uncertainty flags, and slide-ready figures.",
        ],
    )

    add_heading(doc, "8. Results and Discussion Template", 1)
    add_body(
        doc,
        "When approved data are available, the results section should be written around decisions rather than around raw plots. Each well or reservoir interval should report the staged screen outcome, the equations that drove the classification, the evidence against competing explanations, and the uncertainty remaining after core calibration.",
    )
    results_rows = [
        ["Figure 1", "Well-log panel with shaded candidate intervals", "Shows where Rt, Vp/Vs, mu-rho, porosity, and NMR-density evidence overlap."],
        ["Figure 2", "Equation-derived feature crossplots", "Shows hydrate/gas/water/review separation in lambda-rho, mu-rho, Vp/Vs, and saturation-proxy space."],
        ["Figure 3", "Pressure-temperature and stress panel", "Explains whether the interval is admissible and whether stress may explain velocity/stiffness."],
        ["Table 1", "Interval classification table", "Reports gate outcomes, phase evidence, saturation proxy, producibility lane, and uncertainty flags."],
        ["Table 2", "Core-to-log calibration table", "Reports depth-match uncertainty, pressure-core quality, measured saturation/porosity, and confidence weights."],
        ["Discussion", "Failure modes and missing evidence", "Explains intervals that failed because of gas, ice, shale, coal, washout, stress, or missing NMR/DTS."],
    ]
    add_table(doc, ["Result item", "Output", "Purpose"], results_rows, [1.0, 2.75, 3.45], header_fill=NAVY)

    add_heading(doc, "9. Current Evidence Anchors and Missing Pieces", 1)
    add_body(
        doc,
        "The organized source library is strong enough for a scaffold. The remaining gap is not the absence of equations; it is field calibration. The eventual approved-data workflow still needs local clean-sand/shale anchors, matrix-density assumptions, water salinity/resistivity, temperature and pressure calibration, core depth matching, and well-grouped validation labels.",
    )
    add_bullets(
        doc,
        [
            "Mount Elbert public studies support multi-log saturation comparison using NMR, P- and S-wave velocity, resistivity, and core context.",
            "Hydrate-01 public work supports acoustic/resistivity/NMR comparison and shows high hydrate occupancy in target reservoir sands.",
            "Eileen Trend public work supports the idea that structural-stratigraphic segmentation and partial fill must constrain interpretation.",
            "North Slope reservoir-quality and overburden sources support treating rock type and stress as contextual controls rather than hydrate labels.",
        ],
    )

    add_heading(doc, "References and Source Anchors", 1)
    references = [
        "Lee, M.W., and Collett, T.S. (2011). In-situ gas hydrate saturation estimated from various well logs at the Mount Elbert Gas Hydrate Stratigraphic Test Well, Alaska North Slope. Marine and Petroleum Geology, 28(2), 439-449. DOI: 10.1016/j.marpetgeo.2009.06.007.",
        "Haines, S.S., Collett, T., Yoneda, J., Shimoda, N., Boswell, R., and Okinaka, N. (2022). Gas hydrate saturation estimates, gas hydrate occurrence, and reservoir characteristics based on well log data from the Hydrate-01 stratigraphic test well, Alaska North Slope. Energy & Fuels, 36(6), 3040-3050. DOI: 10.1021/acs.energyfuels.1c04100.",
        "Haines, S.S., Collett, T., Boswell, R., Lim, T., Okinaka, N., Suzuki, K., and Fujimoto, A. USGS summary page: Gas hydrate saturation estimation from acoustic log data in the 2018 Alaska North Slope Hydrate-01 stratigraphic test well.",
        "Zyrianova, M., Collett, T., and Boswell, R. (2024). Characterization of the structural-stratigraphic and reservoir controls on the occurrence of gas hydrates in the Eileen Gas Hydrate Trend, Alaska North Slope.",
        "NETL / DOE Mount Elbert project page. The Alaska North Slope Stratigraphic Test Well: public summary of Mount Elbert well-log data types.",
        "Organized local source library: docs/source_library_index/source_index.md and source_manifest.csv. Use these as a citation checklist, not as final provenance.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        for run in p.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(9)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    add_heading(doc, "Appendix A. Sweet-Spot Review Lanes", 1)
    add_body(
        doc,
        "These lanes are designed for review and presentation. They should not be treated as universal field thresholds. The real-data project must recalibrate them against local logs, core observations, reservoir compartment, and uncertainty.",
    )
    sweet_rows = [
        [row["Planning band"], row["Synthetic saturation-proxy range"], row["Use"], row["Required confirmation"]]
        for row in sweet_spot_guide
    ]
    add_table(doc, ["Planning band", "Proxy range", "Use", "Required confirmation"], sweet_rows, [1.75, 1.2, 2.4, 2.6], header_fill=TEAL)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
