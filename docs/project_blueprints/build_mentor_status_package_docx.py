from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "MENTOR_PROJECT_STATUS_PACKAGE_V5_WORKFLOW_2026-06-15.md"
OUT_DOCX = (
    ROOT
    / "docs"
    / "project_blueprints"
    / "North_Slope_Gas_Hydrate_Mentor_Status_Package_V5_Workflow_2026-06-15.docx"
)


def apply_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [
        ("Title", 22),
        ("Heading 1", 16),
        ("Heading 2", 13),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True

    for section in document.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_metadata(document: Document) -> None:
    props = document.core_properties
    props.title = "North Slope Gas Hydrate Mentor Status Package: V5 Workflow"
    props.subject = "Public-safe mentor update around the V5 workflow diagrams"
    props.author = "North Slope Gas Hydrates project"
    props.keywords = "North Slope, gas hydrates, mentor update, stability screen, ML workflow"


def add_markdown(document: Document, source: str) -> None:
    paragraph_lines: list[str] = []
    list_style: str | None = None
    list_text: str | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            document.add_paragraph(" ".join(paragraph_lines))
            paragraph_lines = []

    def flush_list_item() -> None:
        nonlocal list_style, list_text
        if list_text and list_style:
            document.add_paragraph(list_text, style=list_style)
        list_style = None
        list_text = None

    for raw_line in source.splitlines():
        line = raw_line.strip()
        if not line:
            flush_list_item()
            flush_paragraph()
            continue
        if line.startswith("# "):
            flush_list_item()
            flush_paragraph()
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            flush_list_item()
            flush_paragraph()
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            flush_list_item()
            flush_paragraph()
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            flush_list_item()
            flush_paragraph()
            list_style = "List Bullet"
            list_text = line[2:].strip()
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            flush_list_item()
            flush_paragraph()
            list_style = "List Number"
            list_text = line.split(". ", 1)[1].strip()
        elif raw_line.startswith(("  ", "\t")) and list_text:
            list_text = f"{list_text} {line}"
        else:
            flush_list_item()
            paragraph_lines.append(line)

    flush_list_item()
    flush_paragraph()


def build_docx() -> Path:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)

    document = Document()
    apply_document_style(document)
    add_metadata(document)
    add_markdown(document, SOURCE_MD.read_text(encoding="utf-8"))

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)
    return OUT_DOCX


def main() -> None:
    path = build_docx()
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
